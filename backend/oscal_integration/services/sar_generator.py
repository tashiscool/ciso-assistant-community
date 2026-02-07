"""
Security Assessment Report (SAR) Generator Service

Generates Security Assessment Reports in DOCX format from CISO Assistant
compliance assessment data. Follows NIST SP 800-53A / FedRAMP SAR structure.
"""

import json
import logging
import uuid
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO

from .oscal_exporter import OSCALExporter

logger = logging.getLogger(__name__)


@dataclass
class SARSection:
    """Represents a section of the SAR document"""
    title: str
    content: str
    level: int = 1
    subsections: Optional[List['SARSection']] = None


class SARGenerator:
    """
    Generates Security Assessment Report (SAR) documents.

    Sections:
    - Cover Page
    - Executive Summary
    - Assessment Scope
    - Methodology
    - Findings Summary
    - Detailed Findings
    - Risk Summary
    - Recommendations
    - Appendices
    """

    RISK_LEVELS = {
        'very_high': {'label': 'Very High', 'color': '8B0000'},
        'high': {'label': 'High', 'color': 'FF0000'},
        'moderate': {'label': 'Moderate', 'color': 'FFA500'},
        'low': {'label': 'Low', 'color': 'FFD700'},
        'very_low': {'label': 'Very Low', 'color': '008000'},
    }

    FINDING_STATUSES = ['open', 'closed', 'in_progress', 'risk_accepted']

    def __init__(self):
        """Initialize SAR generator"""
        self.exporter = OSCALExporter()

    def generate(self, assessment_id: str, format: str = 'docx', options: Optional[Dict] = None) -> bytes:
        """
        Generate a Security Assessment Report.

        Args:
            assessment_id: UUID of the compliance assessment
            format: Output format ('docx', 'pdf', 'oscal_json')
            options: Generation options (include_appendices, include_evidence, etc.)

        Returns:
            Document content as bytes
        """
        options = options or {}
        assessment_data = self._get_assessment_data(assessment_id)

        if format == 'docx':
            return self._generate_docx(assessment_data, options)
        elif format == 'pdf':
            return self._generate_pdf_placeholder(assessment_data, options)
        elif format == 'oscal_json':
            return self._generate_oscal_json(assessment_data, options)
        else:
            raise ValueError(f"Unsupported SAR format: {format}")

    def _get_assessment_data(self, assessment_id: str) -> Dict[str, Any]:
        """
        Retrieve assessment data for SAR generation.

        Queries compliance assessment, requirement assessments, findings,
        and associated risk data from the database.
        """
        try:
            from core.models import ComplianceAssessment, RequirementAssessment
            from core.models import RiskScenario, AppliedControl

            assessment = ComplianceAssessment.objects.get(id=assessment_id)

            requirement_assessments = RequirementAssessment.objects.filter(
                compliance_assessment=assessment
            ).select_related('requirement')

            # Build assessment data
            total_requirements = requirement_assessments.count()
            compliant = requirement_assessments.filter(result='compliant').count()
            partially_compliant = requirement_assessments.filter(result='partially_compliant').count()
            non_compliant = requirement_assessments.filter(result='non_compliant').count()
            not_assessed = requirement_assessments.filter(result='not_assessed').count()
            not_applicable = requirement_assessments.filter(result='not_applicable').count()

            # Gather findings (non-compliant and partially compliant items)
            findings = []
            for ra in requirement_assessments.filter(
                result__in=['non_compliant', 'partially_compliant']
            ):
                findings.append({
                    'requirement_id': str(ra.requirement.ref_id) if ra.requirement else 'N/A',
                    'requirement_name': str(ra.requirement) if ra.requirement else 'Unknown',
                    'result': ra.result,
                    'observation': ra.observation or '',
                    'risk_level': self._assess_finding_risk(ra),
                })

            return {
                'assessment_id': str(assessment_id),
                'assessment_name': str(assessment),
                'project_name': str(assessment.project) if assessment.project else 'N/A',
                'framework_name': str(assessment.framework) if assessment.framework else 'N/A',
                'status': assessment.status if hasattr(assessment, 'status') else 'in_progress',
                'created_at': assessment.created_at.isoformat() if hasattr(assessment, 'created_at') else datetime.now().isoformat(),
                'summary': {
                    'total_requirements': total_requirements,
                    'compliant': compliant,
                    'partially_compliant': partially_compliant,
                    'non_compliant': non_compliant,
                    'not_assessed': not_assessed,
                    'not_applicable': not_applicable,
                    'compliance_percentage': round((compliant / total_requirements * 100), 1) if total_requirements > 0 else 0,
                },
                'findings': findings,
            }

        except Exception as e:
            logger.warning(f"Could not load assessment data from DB: {e}. Using placeholder data.")
            return self._get_placeholder_data(assessment_id)

    def _assess_finding_risk(self, requirement_assessment) -> str:
        """Determine risk level for a finding based on its requirement assessment"""
        if requirement_assessment.result == 'non_compliant':
            return 'high'
        elif requirement_assessment.result == 'partially_compliant':
            return 'moderate'
        return 'low'

    def _get_placeholder_data(self, assessment_id: str) -> Dict[str, Any]:
        """Return placeholder data when DB queries are not available"""
        return {
            'assessment_id': str(assessment_id),
            'assessment_name': f'Assessment {assessment_id}',
            'project_name': 'Project',
            'framework_name': 'Framework',
            'status': 'in_progress',
            'created_at': datetime.now().isoformat(),
            'summary': {
                'total_requirements': 0,
                'compliant': 0,
                'partially_compliant': 0,
                'non_compliant': 0,
                'not_assessed': 0,
                'not_applicable': 0,
                'compliance_percentage': 0,
            },
            'findings': [],
        }

    def _generate_docx(self, data: Dict[str, Any], options: Dict) -> bytes:
        """Generate SAR as Word document using python-docx"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")

        doc = Document()
        summary = data.get('summary', {})
        findings = data.get('findings', [])

        # -- Cover Page --
        doc.add_paragraph('')
        doc.add_paragraph('')
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('SECURITY ASSESSMENT REPORT')
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(data.get('assessment_name', 'Assessment'))
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"\nProject: {data.get('project_name', 'N/A')}")
        meta.add_run(f"\nFramework: {data.get('framework_name', 'N/A')}")
        meta.add_run(f"\nDate: {datetime.now().strftime('%B %d, %Y')}")
        meta.add_run(f"\nGenerated by CISO Assistant")

        doc.add_page_break()

        # -- Table of Contents placeholder --
        doc.add_heading('Table of Contents', level=1)
        toc_items = [
            '1. Executive Summary',
            '2. Assessment Scope',
            '3. Methodology',
            '4. Findings Summary',
            '5. Detailed Findings',
            '6. Risk Summary',
            '7. Recommendations',
        ]
        if options.get('include_appendices', True):
            toc_items.append('8. Appendices')

        for item in toc_items:
            p = doc.add_paragraph(item)
            p.style = doc.styles['List Bullet']

        doc.add_page_break()

        # -- 1. Executive Summary --
        doc.add_heading('1. Executive Summary', level=1)
        compliance_pct = summary.get('compliance_percentage', 0)
        total = summary.get('total_requirements', 0)
        doc.add_paragraph(
            f"This Security Assessment Report (SAR) documents the findings from the security "
            f"assessment of {data.get('assessment_name', 'the system')}. The assessment evaluated "
            f"{total} requirements against the {data.get('framework_name', 'applicable')} framework."
        )
        doc.add_paragraph(
            f"Overall compliance stands at {compliance_pct}%. "
            f"Of {total} total requirements, {summary.get('compliant', 0)} are fully compliant, "
            f"{summary.get('partially_compliant', 0)} are partially compliant, "
            f"{summary.get('non_compliant', 0)} are non-compliant, "
            f"and {summary.get('not_assessed', 0)} have not yet been assessed."
        )

        # Summary table
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Light Shading Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rows_data = [
            ('Metric', 'Value'),
            ('Total Requirements', str(total)),
            ('Compliant', str(summary.get('compliant', 0))),
            ('Partially Compliant', str(summary.get('partially_compliant', 0))),
            ('Non-Compliant', str(summary.get('non_compliant', 0))),
            ('Not Assessed', str(summary.get('not_assessed', 0))),
            ('Compliance Rate', f"{compliance_pct}%"),
        ]
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            if i == 0:
                for cell in table.rows[i].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        doc.add_paragraph('')

        # -- 2. Assessment Scope --
        doc.add_heading('2. Assessment Scope', level=1)
        doc.add_paragraph(
            f"The assessment covered the {data.get('framework_name', 'applicable')} framework "
            f"as applied to project \"{data.get('project_name', 'N/A')}\". "
            f"All {total} requirements within the framework scope were evaluated."
        )
        doc.add_heading('2.1 In Scope', level=2)
        doc.add_paragraph(
            f"All requirements defined by the {data.get('framework_name', '')} framework "
            f"applicable to the assessed system."
        )
        doc.add_heading('2.2 Out of Scope', level=2)
        na_count = summary.get('not_applicable', 0)
        if na_count > 0:
            doc.add_paragraph(
                f"{na_count} requirements were marked as not applicable and excluded from scoring."
            )
        else:
            doc.add_paragraph("No requirements were excluded from the assessment scope.")

        # -- 3. Methodology --
        doc.add_heading('3. Methodology', level=1)
        doc.add_paragraph(
            "The assessment was conducted using a risk-based approach aligned with NIST SP 800-53A "
            "assessment methodology. Each requirement was evaluated through a combination of:"
        )
        methods = [
            'Document review and evidence examination',
            'Configuration verification and testing',
            'Personnel interviews',
            'Observation of operational processes',
        ]
        for method in methods:
            p = doc.add_paragraph(method)
            p.style = doc.styles['List Bullet']

        doc.add_paragraph(
            "Each requirement was assigned one of the following assessment results: "
            "Compliant, Partially Compliant, Non-Compliant, Not Assessed, or Not Applicable."
        )

        # -- 4. Findings Summary --
        doc.add_heading('4. Findings Summary', level=1)
        if findings:
            doc.add_paragraph(
                f"The assessment identified {len(findings)} findings requiring attention. "
                f"The following table summarizes findings by risk level."
            )
            # Count by risk level
            risk_counts = {}
            for f in findings:
                rl = f.get('risk_level', 'moderate')
                risk_counts[rl] = risk_counts.get(rl, 0) + 1

            risk_table = doc.add_table(rows=len(risk_counts) + 1, cols=2)
            risk_table.style = 'Light Shading Accent 1'
            risk_table.rows[0].cells[0].text = 'Risk Level'
            risk_table.rows[0].cells[1].text = 'Count'
            for cell in risk_table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for i, (level, count) in enumerate(sorted(risk_counts.items()), start=1):
                risk_info = self.RISK_LEVELS.get(level, {'label': level.title()})
                risk_table.rows[i].cells[0].text = risk_info.get('label', level)
                risk_table.rows[i].cells[1].text = str(count)
        else:
            doc.add_paragraph("No findings were identified during the assessment.")

        # -- 5. Detailed Findings --
        doc.add_heading('5. Detailed Findings', level=1)
        if findings:
            for idx, finding in enumerate(findings, start=1):
                doc.add_heading(
                    f"5.{idx} {finding.get('requirement_id', 'N/A')} - {finding.get('requirement_name', 'Unknown')}",
                    level=2
                )
                finding_table = doc.add_table(rows=4, cols=2)
                finding_table.style = 'Light Grid Accent 1'
                finding_data = [
                    ('Requirement', finding.get('requirement_id', 'N/A')),
                    ('Assessment Result', finding.get('result', '').replace('_', ' ').title()),
                    ('Risk Level', self.RISK_LEVELS.get(finding.get('risk_level', ''), {}).get('label', finding.get('risk_level', 'N/A'))),
                    ('Observation', finding.get('observation', 'No observation recorded.')),
                ]
                for row_idx, (label, value) in enumerate(finding_data):
                    finding_table.rows[row_idx].cells[0].text = label
                    finding_table.rows[row_idx].cells[1].text = str(value)
                    for paragraph in finding_table.rows[row_idx].cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                doc.add_paragraph('')
        else:
            doc.add_paragraph("No detailed findings to report.")

        # -- 6. Risk Summary --
        doc.add_heading('6. Risk Summary', level=1)
        high_risk = sum(1 for f in findings if f.get('risk_level') in ('high', 'very_high'))
        moderate_risk = sum(1 for f in findings if f.get('risk_level') == 'moderate')
        low_risk = sum(1 for f in findings if f.get('risk_level') in ('low', 'very_low'))
        doc.add_paragraph(
            f"The overall risk posture based on assessment findings:\n"
            f"- High/Very High Risk Items: {high_risk}\n"
            f"- Moderate Risk Items: {moderate_risk}\n"
            f"- Low/Very Low Risk Items: {low_risk}"
        )
        if high_risk > 0:
            doc.add_paragraph(
                "ATTENTION: High-risk findings require immediate remediation. "
                "A Plan of Action and Milestones (POA&M) should be created for all open findings."
            )

        # -- 7. Recommendations --
        doc.add_heading('7. Recommendations', level=1)
        recommendations = self._generate_recommendations(data)
        for i, rec in enumerate(recommendations, start=1):
            doc.add_paragraph(f"{i}. {rec}")

        # -- 8. Appendices --
        if options.get('include_appendices', True):
            doc.add_page_break()
            doc.add_heading('8. Appendices', level=1)
            doc.add_heading('8.1 Assessment Metadata', level=2)
            doc.add_paragraph(f"Assessment ID: {data.get('assessment_id', 'N/A')}")
            doc.add_paragraph(f"Assessment Status: {data.get('status', 'N/A')}")
            doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Generator: CISO Assistant SAR Generator")

            doc.add_heading('8.2 Abbreviations', level=2)
            abbreviations = [
                ('SAR', 'Security Assessment Report'),
                ('SAP', 'Security Assessment Plan'),
                ('SSP', 'System Security Plan'),
                ('POA&M', 'Plan of Action and Milestones'),
                ('ConMon', 'Continuous Monitoring'),
            ]
            abbr_table = doc.add_table(rows=len(abbreviations) + 1, cols=2)
            abbr_table.style = 'Light Shading Accent 1'
            abbr_table.rows[0].cells[0].text = 'Abbreviation'
            abbr_table.rows[0].cells[1].text = 'Definition'
            for i, (abbr, defn) in enumerate(abbreviations, start=1):
                abbr_table.rows[i].cells[0].text = abbr
                abbr_table.rows[i].cells[1].text = defn

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _generate_pdf_placeholder(self, data: Dict[str, Any], options: Dict) -> bytes:
        """Generate SAR as HTML (PDF placeholder)"""
        summary = data.get('summary', {})
        findings = data.get('findings', [])

        html = f"""<!DOCTYPE html>
<html>
<head>
    <title>Security Assessment Report - {data.get('assessment_name', 'Assessment')}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; color: #333; }}
        h1 {{ color: #1F4E79; border-bottom: 2px solid #1F4E79; padding-bottom: 8px; }}
        h2 {{ color: #2E75B6; }}
        table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #1F4E79; color: white; }}
        .cover {{ text-align: center; padding: 100px 0; }}
        .cover h1 {{ font-size: 32px; border: none; }}
        .metric {{ display: inline-block; padding: 16px; margin: 8px; background: #f0f4f8; border-radius: 8px; text-align: center; min-width: 120px; }}
        .metric .value {{ font-size: 24px; font-weight: bold; color: #1F4E79; }}
        .metric .label {{ font-size: 12px; color: #666; }}
    </style>
</head>
<body>
    <div class="cover">
        <h1>SECURITY ASSESSMENT REPORT</h1>
        <p style="font-size: 18px;">{data.get('assessment_name', 'Assessment')}</p>
        <p>Project: {data.get('project_name', 'N/A')}</p>
        <p>Framework: {data.get('framework_name', 'N/A')}</p>
        <p>Date: {datetime.now().strftime('%B %d, %Y')}</p>
    </div>

    <h1>1. Executive Summary</h1>
    <div>
        <div class="metric"><div class="value">{summary.get('total_requirements', 0)}</div><div class="label">Total Requirements</div></div>
        <div class="metric"><div class="value">{summary.get('compliant', 0)}</div><div class="label">Compliant</div></div>
        <div class="metric"><div class="value">{summary.get('non_compliant', 0)}</div><div class="label">Non-Compliant</div></div>
        <div class="metric"><div class="value">{summary.get('compliance_percentage', 0)}%</div><div class="label">Compliance Rate</div></div>
    </div>

    <h1>2. Findings Summary</h1>
    <table>
        <tr><th>Requirement</th><th>Result</th><th>Risk Level</th><th>Observation</th></tr>
"""
        for finding in findings:
            result_display = finding.get('result', '').replace('_', ' ').title()
            risk_display = self.RISK_LEVELS.get(finding.get('risk_level', ''), {}).get('label', finding.get('risk_level', ''))
            html += f"        <tr><td>{finding.get('requirement_id', 'N/A')}</td><td>{result_display}</td><td>{risk_display}</td><td>{finding.get('observation', '')}</td></tr>\n"

        html += """    </table>

    <p style="text-align: center; color: #999; margin-top: 40px;">
        Generated by CISO Assistant | This HTML document can be converted to PDF using a browser print function.
    </p>
</body>
</html>"""
        return html.encode('utf-8')

    def _generate_oscal_json(self, data: Dict[str, Any], options: Dict) -> bytes:
        """Generate SAR in OSCAL Assessment Results format"""
        findings = data.get('findings', [])
        summary = data.get('summary', {})

        oscal_ar = {
            "assessment-results": {
                "uuid": str(uuid.uuid4()),
                "metadata": {
                    "title": f"Security Assessment Report - {data.get('assessment_name', 'Assessment')}",
                    "last-modified": datetime.now().isoformat(),
                    "version": "1.0",
                    "oscal-version": "1.1.2",
                    "roles": [
                        {"id": "assessor", "title": "Security Assessor"},
                        {"id": "system-owner", "title": "System Owner"},
                    ],
                },
                "import-ap": {
                    "href": f"#assessment-plan-{data.get('assessment_id', '')}",
                },
                "results": [
                    {
                        "uuid": str(uuid.uuid4()),
                        "title": f"Assessment Results for {data.get('assessment_name', '')}",
                        "description": f"Results from assessment of {data.get('framework_name', '')} framework",
                        "start": data.get('created_at', datetime.now().isoformat()),
                        "end": datetime.now().isoformat(),
                        "findings": [
                            {
                                "uuid": str(uuid.uuid4()),
                                "title": f.get('requirement_name', ''),
                                "description": f.get('observation', ''),
                                "target": {
                                    "type": "objective-id",
                                    "target-id": f.get('requirement_id', ''),
                                    "status": {
                                        "state": "not-satisfied" if f.get('result') == 'non_compliant' else "satisfied",
                                    },
                                },
                                "props": [
                                    {"name": "risk-level", "value": f.get('risk_level', 'moderate')},
                                ],
                            }
                            for f in findings
                        ],
                        "props": [
                            {"name": "total-requirements", "value": str(summary.get('total_requirements', 0))},
                            {"name": "compliant", "value": str(summary.get('compliant', 0))},
                            {"name": "non-compliant", "value": str(summary.get('non_compliant', 0))},
                            {"name": "compliance-percentage", "value": str(summary.get('compliance_percentage', 0))},
                        ],
                    }
                ],
            }
        }

        return json.dumps(oscal_ar, indent=2, default=str).encode('utf-8')

    def _generate_recommendations(self, data: Dict[str, Any]) -> List[str]:
        """Generate recommendations based on assessment findings"""
        recommendations = []
        summary = data.get('summary', {})
        findings = data.get('findings', [])

        compliance_pct = summary.get('compliance_percentage', 0)
        non_compliant = summary.get('non_compliant', 0)
        not_assessed = summary.get('not_assessed', 0)
        high_risk = sum(1 for f in findings if f.get('risk_level') in ('high', 'very_high'))

        if high_risk > 0:
            recommendations.append(
                f"Immediately address {high_risk} high-risk finding(s). Develop a POA&M "
                f"with target remediation dates within 30 days."
            )

        if non_compliant > 0:
            recommendations.append(
                f"Create remediation plans for all {non_compliant} non-compliant requirements. "
                f"Prioritize based on risk level and potential impact."
            )

        if not_assessed > 0:
            recommendations.append(
                f"Complete assessment of {not_assessed} remaining requirements to achieve "
                f"full assessment coverage."
            )

        if compliance_pct < 80:
            recommendations.append(
                "Overall compliance is below 80%. Consider a focused remediation sprint "
                "to bring compliance above the acceptable threshold."
            )
        elif compliance_pct < 95:
            recommendations.append(
                "Continue remediation efforts to achieve target compliance threshold of 95% or higher."
            )

        recommendations.append(
            "Establish a continuous monitoring program to maintain compliance posture "
            "and detect control degradation over time."
        )

        recommendations.append(
            "Schedule the next assessment cycle to validate remediation effectiveness "
            "and reassess any accepted risks."
        )

        return recommendations
