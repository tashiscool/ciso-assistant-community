"""
Export Builder Service

Configurable export engine that generates compliance documents in multiple
formats. Serves as the central coordinator for all document generation
in CISO Assistant.

Supported document types: SSP, SAP, SAR, POAM, RISK_REGISTER, CONMON_REPORT
Supported formats: DOCX, XLSX, PDF (HTML placeholder), OSCAL_JSON, OSCAL_YAML, CSV
"""

import csv
import json
import logging
import uuid
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
from datetime import datetime, date
from io import BytesIO, StringIO
from enum import Enum

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    SSP = 'ssp'
    SAP = 'sap'
    SAR = 'sar'
    POAM = 'poam'
    RISK_REGISTER = 'risk_register'
    CONMON_REPORT = 'conmon_report'


class OutputFormat(str, Enum):
    DOCX = 'docx'
    XLSX = 'xlsx'
    PDF = 'pdf'
    OSCAL_JSON = 'oscal_json'
    OSCAL_YAML = 'oscal_yaml'
    CSV = 'csv'


# Content type mapping
CONTENT_TYPES = {
    OutputFormat.DOCX: 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    OutputFormat.XLSX: 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    OutputFormat.PDF: 'text/html',  # HTML placeholder for PDF
    OutputFormat.OSCAL_JSON: 'application/json',
    OutputFormat.OSCAL_YAML: 'application/x-yaml',
    OutputFormat.CSV: 'text/csv',
}

# File extensions
FILE_EXTENSIONS = {
    OutputFormat.DOCX: '.docx',
    OutputFormat.XLSX: '.xlsx',
    OutputFormat.PDF: '.html',
    OutputFormat.OSCAL_JSON: '.json',
    OutputFormat.OSCAL_YAML: '.yaml',
    OutputFormat.CSV: '.csv',
}

# Document type metadata
DOCUMENT_TYPE_INFO = {
    DocumentType.SSP: {
        'name': 'System Security Plan',
        'description': 'Comprehensive security plan documenting control implementations',
        'supported_formats': [OutputFormat.DOCX, OutputFormat.OSCAL_JSON, OutputFormat.OSCAL_YAML, OutputFormat.PDF],
        'icon': 'fa-file-shield',
    },
    DocumentType.SAR: {
        'name': 'Security Assessment Report',
        'description': 'Assessment findings, risk analysis, and recommendations',
        'supported_formats': [OutputFormat.DOCX, OutputFormat.OSCAL_JSON, OutputFormat.PDF],
        'icon': 'fa-clipboard-check',
    },
    DocumentType.SAP: {
        'name': 'Security Assessment Plan',
        'description': 'Assessment scope, methodology, schedule, and team composition',
        'supported_formats': [OutputFormat.DOCX, OutputFormat.OSCAL_JSON, OutputFormat.PDF],
        'icon': 'fa-clipboard-list',
    },
    DocumentType.POAM: {
        'name': 'Plan of Action & Milestones',
        'description': 'Tracking of remediation activities and milestones',
        'supported_formats': [OutputFormat.XLSX, OutputFormat.CSV, OutputFormat.OSCAL_JSON],
        'icon': 'fa-list-check',
    },
    DocumentType.RISK_REGISTER: {
        'name': 'Risk Register',
        'description': 'Comprehensive register of identified risks and treatments',
        'supported_formats': [OutputFormat.XLSX, OutputFormat.CSV, OutputFormat.PDF],
        'icon': 'fa-triangle-exclamation',
    },
    DocumentType.CONMON_REPORT: {
        'name': 'Continuous Monitoring Report',
        'description': 'Periodic monitoring status, metrics, and trend analysis',
        'supported_formats': [OutputFormat.DOCX, OutputFormat.PDF, OutputFormat.XLSX],
        'icon': 'fa-chart-line',
    },
}


@dataclass
class ExportResult:
    """Result of a document export operation"""
    success: bool
    content_bytes: bytes
    filename: str
    content_type: str
    errors: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


class ExportBuilder:
    """
    Central export engine for generating compliance documents.

    Delegates to specialized generators (SSPGenerator, SARGenerator, SAPGenerator)
    and provides unified document generation for all supported types and formats.
    """

    def __init__(self):
        """Initialize export builder with lazy-loaded generators"""
        self._ssp_generator = None
        self._sar_generator = None
        self._sap_generator = None

    @property
    def ssp_generator(self):
        if self._ssp_generator is None:
            from .ssp_generator import SSPGenerator
            self._ssp_generator = SSPGenerator()
        return self._ssp_generator

    @property
    def sar_generator(self):
        if self._sar_generator is None:
            from .sar_generator import SARGenerator
            self._sar_generator = SARGenerator()
        return self._sar_generator

    @property
    def sap_generator(self):
        if self._sap_generator is None:
            from .sap_generator import SAPGenerator
            self._sap_generator = SAPGenerator()
        return self._sap_generator

    def generate_document(
        self,
        doc_type: str,
        format: str,
        assessment_id: Optional[str] = None,
        system_id: Optional[str] = None,
        options: Optional[Dict[str, Any]] = None,
    ) -> ExportResult:
        """
        Generate a compliance document.

        Args:
            doc_type: Document type (ssp, sar, sap, poam, risk_register, conmon_report)
            format: Output format (docx, xlsx, pdf, oscal_json, oscal_yaml, csv)
            assessment_id: UUID of the compliance assessment (for SSP, SAR, SAP)
            system_id: UUID of the system (for risk register, conmon)
            options: Additional generation options

        Returns:
            ExportResult with generated document content
        """
        options = options or {}

        try:
            doc_type_enum = DocumentType(doc_type)
            format_enum = OutputFormat(format)
        except ValueError as e:
            return ExportResult(
                success=False,
                content_bytes=b'',
                filename='',
                content_type='',
                errors=[f"Invalid document type or format: {e}"],
            )

        # Validate format is supported for this document type
        type_info = DOCUMENT_TYPE_INFO.get(doc_type_enum, {})
        supported_formats = type_info.get('supported_formats', [])
        if format_enum not in supported_formats:
            return ExportResult(
                success=False,
                content_bytes=b'',
                filename='',
                content_type='',
                errors=[
                    f"Format '{format}' is not supported for document type '{doc_type}'. "
                    f"Supported formats: {[f.value for f in supported_formats]}"
                ],
            )

        try:
            if doc_type_enum == DocumentType.SSP:
                return self.generate_ssp(assessment_id, format, options)
            elif doc_type_enum == DocumentType.SAR:
                return self.generate_sar(assessment_id, format, options)
            elif doc_type_enum == DocumentType.SAP:
                return self.generate_sap(assessment_id, format, options)
            elif doc_type_enum == DocumentType.POAM:
                return self.generate_poam(assessment_id, format, options)
            elif doc_type_enum == DocumentType.RISK_REGISTER:
                return self.generate_risk_register(format, options.get('filters'))
            elif doc_type_enum == DocumentType.CONMON_REPORT:
                return self.generate_conmon_report(
                    period_start=options.get('period_start'),
                    period_end=options.get('period_end'),
                    format=format,
                )
            else:
                return ExportResult(
                    success=False,
                    content_bytes=b'',
                    filename='',
                    content_type='',
                    errors=[f"Document type '{doc_type}' is not yet implemented"],
                )
        except Exception as e:
            logger.error(f"Error generating {doc_type} document: {e}", exc_info=True)
            return ExportResult(
                success=False,
                content_bytes=b'',
                filename='',
                content_type='',
                errors=[str(e)],
            )

    def generate_ssp(
        self,
        assessment_id: Optional[str],
        format: str = 'docx',
        options: Optional[Dict] = None,
    ) -> ExportResult:
        """
        Generate a System Security Plan.

        For OSCAL formats, delegates to the existing SSPGenerator.
        For DOCX, generates a Word document directly.
        """
        options = options or {}
        format_enum = OutputFormat(format)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        if not assessment_id:
            return ExportResult(
                success=False, content_bytes=b'', filename='', content_type='',
                errors=['assessment_id is required for SSP generation'],
            )

        if format_enum in (OutputFormat.OSCAL_JSON, OutputFormat.OSCAL_YAML):
            # Delegate to existing SSP generator for OSCAL output
            from .oscal_exporter import OSCALExporter
            exporter = OSCALExporter()
            oscal_content = exporter.export_compliance_assessment(assessment_id)

            if format_enum == OutputFormat.OSCAL_YAML:
                try:
                    import yaml
                    json_data = json.loads(oscal_content)
                    oscal_content = yaml.dump(json_data, default_flow_style=False, allow_unicode=True)
                except ImportError:
                    return ExportResult(
                        success=False, content_bytes=b'', filename='', content_type='',
                        errors=['PyYAML is required for YAML export. Install with: pip install pyyaml'],
                    )

            content_bytes = oscal_content.encode('utf-8') if isinstance(oscal_content, str) else oscal_content
            ext = FILE_EXTENSIONS[format_enum]

            return ExportResult(
                success=True,
                content_bytes=content_bytes,
                filename=f"ssp_{assessment_id}_{timestamp}{ext}",
                content_type=CONTENT_TYPES[format_enum],
                metadata={'assessment_id': assessment_id, 'document_type': 'ssp', 'format': format},
            )

        elif format_enum == OutputFormat.DOCX:
            # Generate SSP as Word document using SSPGenerator's appendix A
            baseline = options.get('baseline', 'moderate')
            try:
                content_bytes = self.ssp_generator.generate_appendix_a(assessment_id, baseline)
            except Exception as e:
                # Fall back to a basic DOCX if the trestle transform fails
                logger.warning(f"SSP Appendix A generation failed, building basic DOCX: {e}")
                content_bytes = self._build_basic_ssp_docx(assessment_id, options)

            return ExportResult(
                success=True,
                content_bytes=content_bytes,
                filename=f"ssp_{assessment_id}_{timestamp}.docx",
                content_type=CONTENT_TYPES[OutputFormat.DOCX],
                metadata={'assessment_id': assessment_id, 'document_type': 'ssp', 'format': 'docx'},
            )

        elif format_enum == OutputFormat.PDF:
            content_bytes = self._build_basic_ssp_html(assessment_id, options)
            return ExportResult(
                success=True,
                content_bytes=content_bytes,
                filename=f"ssp_{assessment_id}_{timestamp}.html",
                content_type=CONTENT_TYPES[OutputFormat.PDF],
                metadata={'assessment_id': assessment_id, 'document_type': 'ssp', 'format': 'pdf'},
            )

        return ExportResult(
            success=False, content_bytes=b'', filename='', content_type='',
            errors=[f"Unsupported format '{format}' for SSP"],
        )

    def generate_sar(
        self,
        assessment_id: Optional[str],
        format: str = 'docx',
        options: Optional[Dict] = None,
    ) -> ExportResult:
        """Generate a Security Assessment Report"""
        options = options or {}
        format_enum = OutputFormat(format)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        if not assessment_id:
            return ExportResult(
                success=False, content_bytes=b'', filename='', content_type='',
                errors=['assessment_id is required for SAR generation'],
            )

        content_bytes = self.sar_generator.generate(assessment_id, format, options)
        ext = FILE_EXTENSIONS[format_enum]

        return ExportResult(
            success=True,
            content_bytes=content_bytes,
            filename=f"sar_{assessment_id}_{timestamp}{ext}",
            content_type=CONTENT_TYPES[format_enum],
            metadata={'assessment_id': assessment_id, 'document_type': 'sar', 'format': format},
        )

    def generate_sap(
        self,
        assessment_id: Optional[str],
        format: str = 'docx',
        options: Optional[Dict] = None,
    ) -> ExportResult:
        """Generate a Security Assessment Plan"""
        options = options or {}
        format_enum = OutputFormat(format)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        if not assessment_id:
            return ExportResult(
                success=False, content_bytes=b'', filename='', content_type='',
                errors=['assessment_id is required for SAP generation'],
            )

        content_bytes = self.sap_generator.generate(assessment_id, format, options)
        ext = FILE_EXTENSIONS[format_enum]

        return ExportResult(
            success=True,
            content_bytes=content_bytes,
            filename=f"sap_{assessment_id}_{timestamp}{ext}",
            content_type=CONTENT_TYPES[format_enum],
            metadata={'assessment_id': assessment_id, 'document_type': 'sap', 'format': format},
        )

    def generate_poam(
        self,
        assessment_id: Optional[str] = None,
        format: str = 'xlsx',
        options: Optional[Dict] = None,
    ) -> ExportResult:
        """
        Generate a POA&M export.

        Delegates to the existing POAMExportService for FedRAMP XLSX and CSV,
        falls back to the OSCAL exporter for OSCAL JSON.
        """
        options = options or {}
        format_enum = OutputFormat(format)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        # Get POA&M items from database
        poam_items, system_info = self._get_poam_data(assessment_id)

        if format_enum == OutputFormat.XLSX:
            from poam.services.poam_export import POAMExportService
            service = POAMExportService()
            result = service.export_fedramp_xlsx(poam_items, system_info)
            return ExportResult(
                success=result.success,
                content_bytes=result.content,
                filename=result.filename or f"poam_{timestamp}.xlsx",
                content_type=result.content_type,
                errors=result.errors,
                metadata={'document_type': 'poam', 'format': 'xlsx'},
            )

        elif format_enum == OutputFormat.CSV:
            from poam.services.poam_export import POAMExportService
            service = POAMExportService()
            result = service.export_csv(poam_items)
            return ExportResult(
                success=result.success,
                content_bytes=result.content,
                filename=result.filename or f"poam_{timestamp}.csv",
                content_type=result.content_type,
                errors=result.errors,
                metadata={'document_type': 'poam', 'format': 'csv'},
            )

        elif format_enum == OutputFormat.OSCAL_JSON:
            from poam.services.poam_export import POAMExportService
            service = POAMExportService()
            result = service.export_oscal_poam(poam_items, system_info)
            return ExportResult(
                success=result.success,
                content_bytes=result.content,
                filename=result.filename or f"poam_{timestamp}.json",
                content_type=result.content_type,
                errors=result.errors,
                metadata={'document_type': 'poam', 'format': 'oscal_json'},
            )

        return ExportResult(
            success=False, content_bytes=b'', filename='', content_type='',
            errors=[f"Unsupported format '{format}' for POA&M"],
        )

    def generate_risk_register(
        self,
        format: str = 'xlsx',
        filters: Optional[Dict] = None,
    ) -> ExportResult:
        """
        Generate a Risk Register export.

        Queries risk scenarios from the database and exports them
        as an Excel workbook or CSV.
        """
        filters = filters or {}
        format_enum = OutputFormat(format)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        risk_data = self._get_risk_register_data(filters)

        if format_enum == OutputFormat.XLSX:
            content_bytes = self._build_risk_register_xlsx(risk_data)
            return ExportResult(
                success=True,
                content_bytes=content_bytes,
                filename=f"risk_register_{timestamp}.xlsx",
                content_type=CONTENT_TYPES[OutputFormat.XLSX],
                metadata={'document_type': 'risk_register', 'format': 'xlsx', 'risk_count': len(risk_data)},
            )

        elif format_enum == OutputFormat.CSV:
            content_bytes = self._build_risk_register_csv(risk_data)
            return ExportResult(
                success=True,
                content_bytes=content_bytes,
                filename=f"risk_register_{timestamp}.csv",
                content_type=CONTENT_TYPES[OutputFormat.CSV],
                metadata={'document_type': 'risk_register', 'format': 'csv', 'risk_count': len(risk_data)},
            )

        elif format_enum == OutputFormat.PDF:
            content_bytes = self._build_risk_register_html(risk_data)
            return ExportResult(
                success=True,
                content_bytes=content_bytes,
                filename=f"risk_register_{timestamp}.html",
                content_type=CONTENT_TYPES[OutputFormat.PDF],
                metadata={'document_type': 'risk_register', 'format': 'pdf', 'risk_count': len(risk_data)},
            )

        return ExportResult(
            success=False, content_bytes=b'', filename='', content_type='',
            errors=[f"Unsupported format '{format}' for Risk Register"],
        )

    def generate_conmon_report(
        self,
        period_start: Optional[str] = None,
        period_end: Optional[str] = None,
        format: str = 'docx',
    ) -> ExportResult:
        """
        Generate a Continuous Monitoring Report.

        Summarizes compliance status, control changes, findings,
        and risk trends over a specified period.
        """
        format_enum = OutputFormat(format)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        # Default to last 30 days if no period specified
        if not period_end:
            period_end = datetime.now().strftime('%Y-%m-%d')
        if not period_start:
            end_dt = datetime.strptime(period_end, '%Y-%m-%d')
            period_start = (end_dt - __import__('datetime').timedelta(days=30)).strftime('%Y-%m-%d')

        conmon_data = self._get_conmon_data(period_start, period_end)

        if format_enum == OutputFormat.DOCX:
            content_bytes = self._build_conmon_docx(conmon_data, period_start, period_end)
            return ExportResult(
                success=True,
                content_bytes=content_bytes,
                filename=f"conmon_report_{period_start}_to_{period_end}_{timestamp}.docx",
                content_type=CONTENT_TYPES[OutputFormat.DOCX],
                metadata={'document_type': 'conmon_report', 'format': 'docx', 'period_start': period_start, 'period_end': period_end},
            )

        elif format_enum == OutputFormat.PDF:
            content_bytes = self._build_conmon_html(conmon_data, period_start, period_end)
            return ExportResult(
                success=True,
                content_bytes=content_bytes,
                filename=f"conmon_report_{period_start}_to_{period_end}_{timestamp}.html",
                content_type=CONTENT_TYPES[OutputFormat.PDF],
                metadata={'document_type': 'conmon_report', 'format': 'pdf', 'period_start': period_start, 'period_end': period_end},
            )

        elif format_enum == OutputFormat.XLSX:
            content_bytes = self._build_conmon_xlsx(conmon_data, period_start, period_end)
            return ExportResult(
                success=True,
                content_bytes=content_bytes,
                filename=f"conmon_report_{period_start}_to_{period_end}_{timestamp}.xlsx",
                content_type=CONTENT_TYPES[OutputFormat.XLSX],
                metadata={'document_type': 'conmon_report', 'format': 'xlsx', 'period_start': period_start, 'period_end': period_end},
            )

        return ExportResult(
            success=False, content_bytes=b'', filename='', content_type='',
            errors=[f"Unsupported format '{format}' for ConMon Report"],
        )

    # -------------------------------------------------------------------------
    # Data retrieval helpers
    # -------------------------------------------------------------------------

    def _get_poam_data(self, assessment_id: Optional[str] = None):
        """Retrieve POA&M items from the database"""
        try:
            from poam.models.poam_item import POAMItem

            system_info = {'name': 'System', 'id': str(assessment_id or '')}

            qs = POAMItem.objects.all()
            if assessment_id:
                qs = qs.filter(assessment_id=assessment_id)
                try:
                    from core.models import ComplianceAssessment
                    assessment = ComplianceAssessment.objects.get(id=assessment_id)
                    system_info['name'] = str(assessment)
                except Exception:
                    pass

            poam_items = []
            for item in qs:
                poam_items.append({
                    'id': str(item.id),
                    'weakness_id': item.weakness_id,
                    'title': item.title,
                    'description': item.description or '',
                    'control_id': item.control_id or '',
                    'status': item.status,
                    'risk_level': item.risk_level,
                    'source_type': item.source_type or '',
                    'source_reference': item.source_reference or '',
                    'identified_date': str(item.identified_date) if item.identified_date else '',
                    'estimated_completion_date': str(item.estimated_completion_date) if item.estimated_completion_date else '',
                    'actual_completion_date': str(item.actual_completion_date) if item.actual_completion_date else '',
                    'point_of_contact': item.point_of_contact or '',
                    'resources_required': item.resources_required or '',
                    'remediation_plan': item.remediation_plan or '',
                    'milestones': item.milestones or [],
                    'comments': item.comments or '',
                    'has_deviation': bool(getattr(item, 'deviation_request', None)),
                    'is_overdue': item.is_overdue,
                })

            return poam_items, system_info
        except Exception as e:
            logger.warning(f"Could not load POA&M data: {e}")
            return [], {'name': 'System', 'id': str(assessment_id or '')}

    def _get_risk_register_data(self, filters: Dict) -> List[Dict[str, Any]]:
        """Retrieve risk scenarios for the risk register"""
        try:
            from core.models import RiskScenario
            qs = RiskScenario.objects.all()

            if filters.get('project_id'):
                qs = qs.filter(risk_assessment__project_id=filters['project_id'])
            if filters.get('risk_level'):
                qs = qs.filter(current_level=filters['risk_level'])

            risks = []
            for rs in qs.select_related('risk_assessment')[:500]:
                risks.append({
                    'id': str(rs.id),
                    'name': str(rs),
                    'description': getattr(rs, 'description', ''),
                    'risk_assessment': str(rs.risk_assessment) if rs.risk_assessment else 'N/A',
                    'current_level': getattr(rs, 'current_level', 'N/A'),
                    'residual_level': getattr(rs, 'residual_level', 'N/A'),
                    'treatment': getattr(rs, 'treatment', 'N/A'),
                    'owner': getattr(rs, 'owner', 'N/A'),
                    'status': getattr(rs, 'status', 'open'),
                })
            return risks
        except Exception as e:
            logger.warning(f"Could not load risk register data: {e}")
            return []

    def _get_conmon_data(self, period_start: str, period_end: str) -> Dict[str, Any]:
        """Retrieve continuous monitoring data for the reporting period"""
        try:
            from core.models import ComplianceAssessment, RequirementAssessment

            assessments = ComplianceAssessment.objects.all()[:10]
            total_assessments = assessments.count()

            # Aggregate compliance stats across assessments
            total_requirements = 0
            compliant = 0
            non_compliant = 0

            for assessment in assessments:
                ras = RequirementAssessment.objects.filter(compliance_assessment=assessment)
                total_requirements += ras.count()
                compliant += ras.filter(result='compliant').count()
                non_compliant += ras.filter(result='non_compliant').count()

            return {
                'period_start': period_start,
                'period_end': period_end,
                'total_assessments': total_assessments,
                'total_requirements': total_requirements,
                'compliant': compliant,
                'non_compliant': non_compliant,
                'compliance_rate': round((compliant / total_requirements * 100), 1) if total_requirements > 0 else 0,
            }
        except Exception as e:
            logger.warning(f"Could not load ConMon data: {e}")
            return {
                'period_start': period_start,
                'period_end': period_end,
                'total_assessments': 0,
                'total_requirements': 0,
                'compliant': 0,
                'non_compliant': 0,
                'compliance_rate': 0,
            }

    # -------------------------------------------------------------------------
    # SSP builders (fallback)
    # -------------------------------------------------------------------------

    def _build_basic_ssp_docx(self, assessment_id: str, options: Dict) -> bytes:
        """Build a basic SSP Word document as fallback"""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('SYSTEM SECURITY PLAN')
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        doc.add_paragraph(f"\nAssessment ID: {assessment_id}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Generated by CISO Assistant")

        doc.add_page_break()
        doc.add_heading('1. System Description', level=1)
        doc.add_paragraph('[System description to be completed]')
        doc.add_heading('2. Control Implementation', level=1)
        doc.add_paragraph('[Control implementation details to be completed]')
        doc.add_heading('3. System Architecture', level=1)
        doc.add_paragraph('[System architecture details to be completed]')

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _build_basic_ssp_html(self, assessment_id: str, options: Dict) -> bytes:
        """Build SSP as HTML for PDF placeholder"""
        html = f"""<!DOCTYPE html>
<html>
<head><title>System Security Plan</title>
<style>body {{ font-family: Arial, sans-serif; margin: 40px; }} h1 {{ color: #1F4E79; }}</style>
</head>
<body>
<h1 style="text-align:center;">SYSTEM SECURITY PLAN</h1>
<p style="text-align:center;">Assessment ID: {assessment_id}<br>Generated: {datetime.now().strftime('%B %d, %Y')}</p>
<h1>1. System Description</h1><p>[To be completed]</p>
<h1>2. Control Implementation</h1><p>[To be completed]</p>
<p style="text-align:center; color:#999;">Generated by CISO Assistant</p>
</body></html>"""
        return html.encode('utf-8')

    # -------------------------------------------------------------------------
    # Risk Register builders
    # -------------------------------------------------------------------------

    def _build_risk_register_xlsx(self, risk_data: List[Dict]) -> bytes:
        """Build risk register as Excel workbook"""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = "Risk Register"

        header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=10)
        border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin'),
        )

        headers = ['Risk ID', 'Name', 'Description', 'Risk Assessment', 'Current Level',
                    'Residual Level', 'Treatment', 'Owner', 'Status']

        for col_idx, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col_idx)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = border
            ws.column_dimensions[get_column_letter(col_idx)].width = 18

        for row_idx, risk in enumerate(risk_data, start=2):
            row_values = [
                risk.get('id', ''), risk.get('name', ''), risk.get('description', ''),
                risk.get('risk_assessment', ''), risk.get('current_level', ''),
                risk.get('residual_level', ''), risk.get('treatment', ''),
                str(risk.get('owner', '')), risk.get('status', ''),
            ]
            for col_idx, value in enumerate(row_values, start=1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.value = value
                cell.border = border
                cell.alignment = Alignment(vertical='top', wrap_text=True)

        ws.freeze_panes = 'A2'

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _build_risk_register_csv(self, risk_data: List[Dict]) -> bytes:
        """Build risk register as CSV"""
        output = StringIO()
        fieldnames = ['id', 'name', 'description', 'risk_assessment', 'current_level',
                       'residual_level', 'treatment', 'owner', 'status']
        writer = csv.DictWriter(output, fieldnames=fieldnames, extrasaction='ignore')
        writer.writeheader()
        for risk in risk_data:
            writer.writerow(risk)
        return output.getvalue().encode('utf-8')

    def _build_risk_register_html(self, risk_data: List[Dict]) -> bytes:
        """Build risk register as HTML for PDF placeholder"""
        html = """<!DOCTYPE html>
<html><head><title>Risk Register</title>
<style>
body { font-family: Arial, sans-serif; margin: 40px; }
h1 { color: #1F4E79; }
table { border-collapse: collapse; width: 100%; }
th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
th { background-color: #1F4E79; color: white; }
tr:nth-child(even) { background-color: #f2f2f2; }
</style></head><body>
<h1>Risk Register</h1>
<p>Generated: """ + datetime.now().strftime('%B %d, %Y') + f"""</p>
<p>Total Risks: {len(risk_data)}</p>
<table>
<tr><th>ID</th><th>Name</th><th>Current Level</th><th>Residual Level</th><th>Treatment</th><th>Status</th></tr>
"""
        for risk in risk_data:
            html += f"<tr><td>{risk.get('id', '')[:8]}</td><td>{risk.get('name', '')}</td>"
            html += f"<td>{risk.get('current_level', '')}</td><td>{risk.get('residual_level', '')}</td>"
            html += f"<td>{risk.get('treatment', '')}</td><td>{risk.get('status', '')}</td></tr>\n"

        html += """</table>
<p style="text-align:center; color:#999;">Generated by CISO Assistant</p>
</body></html>"""
        return html.encode('utf-8')

    # -------------------------------------------------------------------------
    # ConMon Report builders
    # -------------------------------------------------------------------------

    def _build_conmon_docx(self, data: Dict, period_start: str, period_end: str) -> bytes:
        """Build continuous monitoring report as Word document"""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        # Cover
        doc.add_paragraph('')
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('CONTINUOUS MONITORING REPORT')
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        period = doc.add_paragraph()
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period.add_run(f"\nReporting Period: {period_start} to {period_end}")
        period.add_run(f"\nGenerated: {datetime.now().strftime('%B %d, %Y')}")
        period.add_run(f"\nGenerated by CISO Assistant")

        doc.add_page_break()

        # Executive Summary
        doc.add_heading('1. Executive Summary', level=1)
        compliance_rate = data.get('compliance_rate', 0)
        doc.add_paragraph(
            f"This report covers the continuous monitoring period from {period_start} to {period_end}. "
            f"During this period, {data.get('total_assessments', 0)} assessments were active, "
            f"covering {data.get('total_requirements', 0)} requirements. "
            f"The overall compliance rate is {compliance_rate}%."
        )

        # Compliance Metrics
        doc.add_heading('2. Compliance Metrics', level=1)
        metrics_table = doc.add_table(rows=5, cols=2)
        metrics_table.style = 'Light Shading Accent 1'
        metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        metrics = [
            ('Metric', 'Value'),
            ('Active Assessments', str(data.get('total_assessments', 0))),
            ('Total Requirements Monitored', str(data.get('total_requirements', 0))),
            ('Compliant Requirements', str(data.get('compliant', 0))),
            ('Non-Compliant Requirements', str(data.get('non_compliant', 0))),
        ]
        for i, (label, value) in enumerate(metrics):
            metrics_table.rows[i].cells[0].text = label
            metrics_table.rows[i].cells[1].text = value
            if i == 0:
                for cell in metrics_table.rows[i].cells:
                    for paragraph in cell.paragraphs:
                        for r in paragraph.runs:
                            r.bold = True

        # Status Assessment
        doc.add_heading('3. Status Assessment', level=1)
        if compliance_rate >= 95:
            doc.add_paragraph("Overall system security posture is SATISFACTORY. Compliance rate exceeds 95% threshold.")
        elif compliance_rate >= 80:
            doc.add_paragraph("Overall system security posture is ACCEPTABLE WITH CONCERNS. Compliance rate is between 80-95%. Remediation recommended.")
        else:
            doc.add_paragraph("Overall system security posture is UNSATISFACTORY. Compliance rate is below 80%. Immediate remediation required.")

        # Recommendations
        doc.add_heading('4. Recommendations', level=1)
        if data.get('non_compliant', 0) > 0:
            doc.add_paragraph(f"Address {data.get('non_compliant', 0)} non-compliant requirements through targeted remediation.")
        doc.add_paragraph("Continue monitoring activities per the established schedule.")
        doc.add_paragraph("Update POA&M entries for any newly identified findings.")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _build_conmon_html(self, data: Dict, period_start: str, period_end: str) -> bytes:
        """Build ConMon report as HTML for PDF placeholder"""
        compliance_rate = data.get('compliance_rate', 0)
        status_color = '#28a745' if compliance_rate >= 95 else '#ffc107' if compliance_rate >= 80 else '#dc3545'

        html = f"""<!DOCTYPE html>
<html><head><title>Continuous Monitoring Report</title>
<style>
body {{ font-family: Arial, sans-serif; margin: 40px; }}
h1 {{ color: #1F4E79; border-bottom: 2px solid #1F4E79; padding-bottom: 8px; }}
.metric {{ display: inline-block; padding: 20px; margin: 8px; background: #f0f4f8; border-radius: 8px; text-align: center; min-width: 150px; }}
.metric .value {{ font-size: 28px; font-weight: bold; color: #1F4E79; }}
.metric .label {{ font-size: 12px; color: #666; }}
.status {{ padding: 12px 24px; border-radius: 8px; color: white; background: {status_color}; display: inline-block; font-weight: bold; }}
</style></head><body>
<h1 style="text-align:center; border:none;">CONTINUOUS MONITORING REPORT</h1>
<p style="text-align:center;">Period: {period_start} to {period_end}</p>

<h1>Compliance Metrics</h1>
<div>
    <div class="metric"><div class="value">{data.get('total_assessments', 0)}</div><div class="label">Assessments</div></div>
    <div class="metric"><div class="value">{data.get('total_requirements', 0)}</div><div class="label">Requirements</div></div>
    <div class="metric"><div class="value">{data.get('compliant', 0)}</div><div class="label">Compliant</div></div>
    <div class="metric"><div class="value">{compliance_rate}%</div><div class="label">Compliance Rate</div></div>
</div>

<h1>Status</h1>
<p><span class="status">{'SATISFACTORY' if compliance_rate >= 95 else 'NEEDS ATTENTION' if compliance_rate >= 80 else 'UNSATISFACTORY'}</span></p>

<p style="text-align:center; color:#999; margin-top: 40px;">Generated by CISO Assistant</p>
</body></html>"""
        return html.encode('utf-8')

    def _build_conmon_xlsx(self, data: Dict, period_start: str, period_end: str) -> bytes:
        """Build ConMon report as Excel workbook"""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = Workbook()
        ws = wb.active
        ws.title = "ConMon Report"

        # Title
        ws.merge_cells('A1:D1')
        ws['A1'] = 'Continuous Monitoring Report'
        ws['A1'].font = Font(bold=True, size=16, color='1F4E79')

        ws['A2'] = f'Period: {period_start} to {period_end}'
        ws['A3'] = f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}'

        # Metrics
        header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)

        ws['A5'] = 'Metric'
        ws['B5'] = 'Value'
        ws['A5'].fill = header_fill
        ws['A5'].font = header_font
        ws['B5'].fill = header_fill
        ws['B5'].font = header_font

        metrics = [
            ('Active Assessments', data.get('total_assessments', 0)),
            ('Total Requirements', data.get('total_requirements', 0)),
            ('Compliant', data.get('compliant', 0)),
            ('Non-Compliant', data.get('non_compliant', 0)),
            ('Compliance Rate', f"{data.get('compliance_rate', 0)}%"),
        ]
        for i, (label, value) in enumerate(metrics, start=6):
            ws[f'A{i}'] = label
            ws[f'A{i}'].font = Font(bold=True)
            ws[f'B{i}'] = value

        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 20

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # -------------------------------------------------------------------------
    # Class methods for document type info
    # -------------------------------------------------------------------------

    @staticmethod
    def get_document_types() -> List[Dict[str, Any]]:
        """Get list of available document types with metadata"""
        result = []
        for doc_type, info in DOCUMENT_TYPE_INFO.items():
            result.append({
                'type': doc_type.value,
                'name': info['name'],
                'description': info['description'],
                'icon': info['icon'],
                'supported_formats': [f.value for f in info['supported_formats']],
            })
        return result

    @staticmethod
    def get_format_info(format_key: str) -> Dict[str, str]:
        """Get info about an output format"""
        try:
            fmt = OutputFormat(format_key)
            return {
                'format': fmt.value,
                'content_type': CONTENT_TYPES.get(fmt, ''),
                'extension': FILE_EXTENSIONS.get(fmt, ''),
            }
        except ValueError:
            return {}
