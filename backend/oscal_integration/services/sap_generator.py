"""
Security Assessment Plan (SAP) Generator Service

Generates Security Assessment Plan documents in DOCX format from CISO Assistant
compliance assessment data. Follows NIST SP 800-53A / FedRAMP SAP structure.
"""

import json
import logging
import uuid
from typing import Dict, List, Any, Optional
from datetime import datetime, timedelta
from io import BytesIO

from .oscal_exporter import OSCALExporter

logger = logging.getLogger(__name__)


class SAPGenerator:
    """
    Generates Security Assessment Plan (SAP) documents.

    Sections:
    - Cover Page
    - Assessment Overview
    - Scope
    - Methodology
    - Schedule
    - Assessment Team
    - Rules of Engagement
    - Communication Plan
    """

    ASSESSMENT_METHODS = [
        {
            'name': 'Examine',
            'description': 'Review and analysis of documentation, policies, procedures, and configurations.',
        },
        {
            'name': 'Interview',
            'description': 'Discussions with key personnel to verify implementation and understanding of controls.',
        },
        {
            'name': 'Test',
            'description': 'Hands-on testing and verification of control implementation and effectiveness.',
        },
    ]

    def __init__(self):
        """Initialize SAP generator"""
        self.exporter = OSCALExporter()

    def generate(self, assessment_id: str, format: str = 'docx', options: Optional[Dict] = None) -> bytes:
        """
        Generate a Security Assessment Plan.

        Args:
            assessment_id: UUID of the compliance assessment
            format: Output format ('docx', 'pdf', 'oscal_json')
            options: Generation options

        Returns:
            Document content as bytes
        """
        options = options or {}
        assessment_data = self._get_assessment_data(assessment_id)

        if format == 'docx':
            return self._generate_docx(assessment_data, options)
        elif format == 'pdf':
            return self._generate_pdf_placeholder(assessment_data, options)
        elif format == 'oscal_json':
            return self._generate_oscal_json(assessment_data, options)
        else:
            raise ValueError(f"Unsupported SAP format: {format}")

    def _get_assessment_data(self, assessment_id: str) -> Dict[str, Any]:
        """
        Retrieve assessment data for SAP generation.

        Queries the compliance assessment and its framework/requirements
        from the database.
        """
        try:
            from core.models import ComplianceAssessment, RequirementAssessment

            assessment = ComplianceAssessment.objects.get(id=assessment_id)

            requirement_assessments = RequirementAssessment.objects.filter(
                compliance_assessment=assessment
            ).select_related('requirement')

            total_requirements = requirement_assessments.count()

            # Group requirements by control family / category
            control_families = {}
            for ra in requirement_assessments:
                if ra.requirement:
                    ref_id = str(ra.requirement.ref_id) if ra.requirement.ref_id else 'OTHER'
                    family = ref_id.split('-')[0] if '-' in ref_id else ref_id[:2]
                    if family not in control_families:
                        control_families[family] = []
                    control_families[family].append({
                        'ref_id': ref_id,
                        'name': str(ra.requirement),
                    })

            return {
                'assessment_id': str(assessment_id),
                'assessment_name': str(assessment),
                'project_name': str(assessment.project) if assessment.project else 'N/A',
                'framework_name': str(assessment.framework) if assessment.framework else 'N/A',
                'status': assessment.status if hasattr(assessment, 'status') else 'planned',
                'created_at': assessment.created_at.isoformat() if hasattr(assessment, 'created_at') else datetime.now().isoformat(),
                'total_requirements': total_requirements,
                'control_families': control_families,
            }

        except Exception as e:
            logger.warning(f"Could not load assessment data from DB: {e}. Using placeholder data.")
            return self._get_placeholder_data(assessment_id)

    def _get_placeholder_data(self, assessment_id: str) -> Dict[str, Any]:
        """Return placeholder data when DB queries are not available"""
        return {
            'assessment_id': str(assessment_id),
            'assessment_name': f'Assessment {assessment_id}',
            'project_name': 'Project',
            'framework_name': 'Framework',
            'status': 'planned',
            'created_at': datetime.now().isoformat(),
            'total_requirements': 0,
            'control_families': {},
        }

    def _generate_docx(self, data: Dict[str, Any], options: Dict) -> bytes:
        """Generate SAP as Word document using python-docx"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")

        doc = Document()
        total_reqs = data.get('total_requirements', 0)
        control_families = data.get('control_families', {})
        now = datetime.now()

        # -- Cover Page --
        doc.add_paragraph('')
        doc.add_paragraph('')
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('SECURITY ASSESSMENT PLAN')
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(data.get('assessment_name', 'Assessment'))
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"\nProject: {data.get('project_name', 'N/A')}")
        meta.add_run(f"\nFramework: {data.get('framework_name', 'N/A')}")
        meta.add_run(f"\nDate: {now.strftime('%B %d, %Y')}")
        meta.add_run(f"\nVersion: 1.0")
        meta.add_run(f"\nGenerated by CISO Assistant")

        doc.add_page_break()

        # -- Table of Contents --
        doc.add_heading('Table of Contents', level=1)
        toc_items = [
            '1. Assessment Overview',
            '2. Scope',
            '3. Methodology',
            '4. Schedule',
            '5. Assessment Team',
            '6. Rules of Engagement',
            '7. Communication Plan',
        ]
        for item in toc_items:
            p = doc.add_paragraph(item)
            p.style = doc.styles['List Bullet']

        doc.add_page_break()

        # -- 1. Assessment Overview --
        doc.add_heading('1. Assessment Overview', level=1)

        doc.add_heading('1.1 Purpose', level=2)
        doc.add_paragraph(
            f"This Security Assessment Plan (SAP) defines the plan for conducting a comprehensive "
            f"security assessment of {data.get('assessment_name', 'the system')}. The assessment "
            f"will evaluate the implementation and effectiveness of security controls as defined "
            f"by the {data.get('framework_name', 'applicable')} framework."
        )

        doc.add_heading('1.2 System Information', level=2)
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Light Shading Accent 1'
        info_data = [
            ('Assessment Name', data.get('assessment_name', 'N/A')),
            ('Project', data.get('project_name', 'N/A')),
            ('Framework', data.get('framework_name', 'N/A')),
            ('Total Requirements', str(total_reqs)),
            ('Control Families', str(len(control_families))),
        ]
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
            for paragraph in info_table.rows[i].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        doc.add_heading('1.3 Assessment Objectives', level=2)
        objectives = [
            'Determine if security controls are implemented correctly',
            'Determine if security controls are operating as intended',
            'Determine if security controls are producing the desired outcome',
            'Identify security weaknesses and deficiencies',
            'Provide recommendations for remediation',
        ]
        for obj in objectives:
            p = doc.add_paragraph(obj)
            p.style = doc.styles['List Bullet']

        # -- 2. Scope --
        doc.add_heading('2. Scope', level=1)

        doc.add_heading('2.1 Assessment Boundaries', level=2)
        doc.add_paragraph(
            f"The assessment covers {total_reqs} requirements across "
            f"{len(control_families)} control families within the "
            f"{data.get('framework_name', '')} framework."
        )

        doc.add_heading('2.2 Control Families in Scope', level=2)
        if control_families:
            scope_table = doc.add_table(rows=len(control_families) + 1, cols=3)
            scope_table.style = 'Light Shading Accent 1'
            scope_table.rows[0].cells[0].text = 'Family'
            scope_table.rows[0].cells[1].text = 'Controls'
            scope_table.rows[0].cells[2].text = 'Method'
            for cell in scope_table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for i, (family, controls) in enumerate(sorted(control_families.items()), start=1):
                scope_table.rows[i].cells[0].text = family.upper()
                scope_table.rows[i].cells[1].text = str(len(controls))
                scope_table.rows[i].cells[2].text = 'Examine, Interview, Test'
        else:
            doc.add_paragraph("Control families will be determined during assessment planning.")

        doc.add_heading('2.3 Exclusions', level=2)
        doc.add_paragraph(
            "Any requirements marked as 'Not Applicable' in the compliance assessment "
            "are excluded from the scope of this assessment."
        )

        # -- 3. Methodology --
        doc.add_heading('3. Methodology', level=1)
        doc.add_paragraph(
            "The assessment methodology follows NIST SP 800-53A guidelines using a "
            "combination of examination, interview, and testing procedures."
        )

        for method in self.ASSESSMENT_METHODS:
            doc.add_heading(f"3.{self.ASSESSMENT_METHODS.index(method) + 1} {method['name']}", level=2)
            doc.add_paragraph(method['description'])

        doc.add_heading('3.4 Assessment Rating Scale', level=2)
        rating_table = doc.add_table(rows=6, cols=2)
        rating_table.style = 'Light Shading Accent 1'
        rating_table.rows[0].cells[0].text = 'Rating'
        rating_table.rows[0].cells[1].text = 'Description'
        for cell in rating_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        ratings = [
            ('Compliant', 'Control is fully implemented and operating effectively'),
            ('Partially Compliant', 'Control is partially implemented or has deficiencies'),
            ('Non-Compliant', 'Control is not implemented or fundamentally flawed'),
            ('Not Assessed', 'Control has not yet been evaluated'),
            ('Not Applicable', 'Control is not relevant to the system'),
        ]
        for i, (rating, desc) in enumerate(ratings, start=1):
            rating_table.rows[i].cells[0].text = rating
            rating_table.rows[i].cells[1].text = desc

        # -- 4. Schedule --
        doc.add_heading('4. Schedule', level=1)
        doc.add_paragraph(
            "The following schedule outlines the planned timeline for the security assessment."
        )

        # Generate a sample schedule based on scope
        phases = self._generate_schedule(total_reqs, now)
        schedule_table = doc.add_table(rows=len(phases) + 1, cols=4)
        schedule_table.style = 'Light Shading Accent 1'
        schedule_table.rows[0].cells[0].text = 'Phase'
        schedule_table.rows[0].cells[1].text = 'Activity'
        schedule_table.rows[0].cells[2].text = 'Start Date'
        schedule_table.rows[0].cells[3].text = 'End Date'
        for cell in schedule_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for i, phase in enumerate(phases, start=1):
            schedule_table.rows[i].cells[0].text = phase['phase']
            schedule_table.rows[i].cells[1].text = phase['activity']
            schedule_table.rows[i].cells[2].text = phase['start']
            schedule_table.rows[i].cells[3].text = phase['end']

        # -- 5. Assessment Team --
        doc.add_heading('5. Assessment Team', level=1)
        doc.add_paragraph(
            "The assessment team consists of qualified security professionals "
            "with experience in the applicable framework and assessment methodologies."
        )

        team_table = doc.add_table(rows=5, cols=3)
        team_table.style = 'Light Shading Accent 1'
        team_table.rows[0].cells[0].text = 'Role'
        team_table.rows[0].cells[1].text = 'Name'
        team_table.rows[0].cells[2].text = 'Responsibilities'
        for cell in team_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        team_roles = [
            ('Lead Assessor', '[To be assigned]', 'Overall assessment leadership and quality assurance'),
            ('Security Assessor', '[To be assigned]', 'Control testing and evidence review'),
            ('Technical Assessor', '[To be assigned]', 'Technical control verification and testing'),
            ('Assessment Coordinator', '[To be assigned]', 'Scheduling, logistics, and communication'),
        ]
        for i, (role, name, resp) in enumerate(team_roles, start=1):
            team_table.rows[i].cells[0].text = role
            team_table.rows[i].cells[1].text = options.get(f'team_{role.lower().replace(" ", "_")}', name)
            team_table.rows[i].cells[2].text = resp

        # -- 6. Rules of Engagement --
        doc.add_heading('6. Rules of Engagement', level=1)

        doc.add_heading('6.1 Assessment Constraints', level=2)
        constraints = [
            'All testing must be coordinated with the system owner prior to execution',
            'No destructive testing or denial-of-service testing is authorized',
            'All assessment activities must be conducted during approved maintenance windows',
            'Sensitive data encountered during assessment must be handled in accordance with data classification policies',
            'Any critical vulnerabilities discovered must be reported immediately to the system owner',
        ]
        for constraint in constraints:
            p = doc.add_paragraph(constraint)
            p.style = doc.styles['List Bullet']

        doc.add_heading('6.2 Access Requirements', level=2)
        doc.add_paragraph(
            "The assessment team requires the following access to conduct the assessment:"
        )
        access_items = [
            'Read access to system documentation and policies',
            'Interview access to key personnel (system administrators, security staff, management)',
            'Read-only access to system configurations and logs',
            'Access to evidence repositories and artifact storage',
        ]
        for item in access_items:
            p = doc.add_paragraph(item)
            p.style = doc.styles['List Bullet']

        doc.add_heading('6.3 Data Handling', level=2)
        doc.add_paragraph(
            "All assessment data, findings, and evidence will be classified and handled "
            "in accordance with the organization's data classification policy. Assessment "
            "results will be shared only with authorized personnel."
        )

        # -- 7. Communication Plan --
        doc.add_heading('7. Communication Plan', level=1)

        doc.add_heading('7.1 Reporting Schedule', level=2)
        comm_items = [
            'Daily status updates to assessment coordinator',
            'Weekly progress reports to system owner and ISSO',
            'Immediate notification of critical findings to system owner',
            'Draft SAR delivery within 5 business days of assessment completion',
            'Final SAR delivery within 10 business days of assessment completion',
        ]
        for item in comm_items:
            p = doc.add_paragraph(item)
            p.style = doc.styles['List Bullet']

        doc.add_heading('7.2 Escalation Procedures', level=2)
        doc.add_paragraph(
            "Issues or disputes arising during the assessment will be escalated as follows:"
        )
        escalation = [
            'Level 1: Lead Assessor and System Owner',
            'Level 2: Assessment Program Manager and ISSO',
            'Level 3: Authorizing Official',
        ]
        for item in escalation:
            p = doc.add_paragraph(item)
            p.style = doc.styles['List Number']

        doc.add_heading('7.3 Points of Contact', level=2)
        poc_table = doc.add_table(rows=4, cols=3)
        poc_table.style = 'Light Shading Accent 1'
        poc_table.rows[0].cells[0].text = 'Role'
        poc_table.rows[0].cells[1].text = 'Name'
        poc_table.rows[0].cells[2].text = 'Contact'
        for cell in poc_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        poc_roles = [
            ('System Owner', '[To be assigned]', '[Email/Phone]'),
            ('ISSO', '[To be assigned]', '[Email/Phone]'),
            ('Lead Assessor', '[To be assigned]', '[Email/Phone]'),
        ]
        for i, (role, name, contact) in enumerate(poc_roles, start=1):
            poc_table.rows[i].cells[0].text = role
            poc_table.rows[i].cells[1].text = name
            poc_table.rows[i].cells[2].text = contact

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _generate_pdf_placeholder(self, data: Dict[str, Any], options: Dict) -> bytes:
        """Generate SAP as HTML (PDF placeholder)"""
        total_reqs = data.get('total_requirements', 0)
        control_families = data.get('control_families', {})

        html = f"""<!DOCTYPE html>
<html>
<head>
    <title>Security Assessment Plan - {data.get('assessment_name', 'Assessment')}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; color: #333; }}
        h1 {{ color: #1F4E79; border-bottom: 2px solid #1F4E79; padding-bottom: 8px; }}
        h2 {{ color: #2E75B6; }}
        table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #1F4E79; color: white; }}
        .cover {{ text-align: center; padding: 100px 0; }}
        .cover h1 {{ font-size: 32px; border: none; }}
    </style>
</head>
<body>
    <div class="cover">
        <h1>SECURITY ASSESSMENT PLAN</h1>
        <p style="font-size: 18px;">{data.get('assessment_name', 'Assessment')}</p>
        <p>Project: {data.get('project_name', 'N/A')}</p>
        <p>Framework: {data.get('framework_name', 'N/A')}</p>
        <p>Date: {datetime.now().strftime('%B %d, %Y')}</p>
    </div>

    <h1>1. Assessment Overview</h1>
    <p>This SAP covers the assessment of {total_reqs} requirements across {len(control_families)} control families.</p>

    <h1>2. Scope</h1>
    <table>
        <tr><th>Family</th><th>Controls</th></tr>
"""
        for family, controls in sorted(control_families.items()):
            html += f"        <tr><td>{family.upper()}</td><td>{len(controls)}</td></tr>\n"

        html += """    </table>

    <h1>3. Methodology</h1>
    <p>Assessment follows NIST SP 800-53A using Examine, Interview, and Test methods.</p>

    <p style="text-align: center; color: #999; margin-top: 40px;">
        Generated by CISO Assistant | This HTML document can be converted to PDF using a browser print function.
    </p>
</body>
</html>"""
        return html.encode('utf-8')

    def _generate_oscal_json(self, data: Dict[str, Any], options: Dict) -> bytes:
        """Generate SAP in OSCAL Assessment Plan format"""
        control_families = data.get('control_families', {})

        # Build control selection
        controls_selected = []
        for family, controls in control_families.items():
            for ctrl in controls:
                controls_selected.append({
                    "control-id": ctrl.get('ref_id', ''),
                })

        oscal_ap = {
            "assessment-plan": {
                "uuid": str(uuid.uuid4()),
                "metadata": {
                    "title": f"Security Assessment Plan - {data.get('assessment_name', 'Assessment')}",
                    "last-modified": datetime.now().isoformat(),
                    "version": "1.0",
                    "oscal-version": "1.1.2",
                    "roles": [
                        {"id": "assessor", "title": "Security Assessor"},
                        {"id": "lead-assessor", "title": "Lead Assessor"},
                        {"id": "system-owner", "title": "System Owner"},
                    ],
                },
                "import-ssp": {
                    "href": f"#ssp-{data.get('assessment_id', '')}",
                },
                "local-definitions": {
                    "activities": [
                        {
                            "uuid": str(uuid.uuid4()),
                            "title": method['name'],
                            "description": method['description'],
                        }
                        for method in self.ASSESSMENT_METHODS
                    ],
                },
                "reviewed-controls": {
                    "control-selections": [
                        {
                            "include-controls": controls_selected[:50],  # Limit for readability
                        }
                    ],
                },
                "assessment-subjects": [
                    {
                        "type": "component",
                        "description": f"Assessment of {data.get('assessment_name', 'system')}",
                        "include-all": {},
                    }
                ],
            }
        }

        return json.dumps(oscal_ap, indent=2, default=str).encode('utf-8')

    def _generate_schedule(self, total_requirements: int, start_date: datetime) -> List[Dict[str, str]]:
        """Generate a schedule based on scope size"""
        # Estimate durations based on requirement count
        if total_requirements < 50:
            prep_days, exec_days, report_days = 3, 5, 3
        elif total_requirements < 150:
            prep_days, exec_days, report_days = 5, 10, 5
        elif total_requirements < 300:
            prep_days, exec_days, report_days = 7, 15, 7
        else:
            prep_days, exec_days, report_days = 10, 20, 10

        phases = []
        current = start_date

        # Phase 1: Preparation
        phase_end = current + timedelta(days=prep_days)
        phases.append({
            'phase': '1',
            'activity': 'Assessment Preparation & Planning',
            'start': current.strftime('%Y-%m-%d'),
            'end': phase_end.strftime('%Y-%m-%d'),
        })
        current = phase_end + timedelta(days=1)

        # Phase 2: Document Review
        phase_end = current + timedelta(days=max(3, prep_days - 2))
        phases.append({
            'phase': '2',
            'activity': 'Document Review & Evidence Collection',
            'start': current.strftime('%Y-%m-%d'),
            'end': phase_end.strftime('%Y-%m-%d'),
        })
        current = phase_end + timedelta(days=1)

        # Phase 3: Assessment Execution
        phase_end = current + timedelta(days=exec_days)
        phases.append({
            'phase': '3',
            'activity': 'Assessment Execution (Examine, Interview, Test)',
            'start': current.strftime('%Y-%m-%d'),
            'end': phase_end.strftime('%Y-%m-%d'),
        })
        current = phase_end + timedelta(days=1)

        # Phase 4: Analysis
        phase_end = current + timedelta(days=max(3, report_days - 2))
        phases.append({
            'phase': '4',
            'activity': 'Findings Analysis & Risk Assessment',
            'start': current.strftime('%Y-%m-%d'),
            'end': phase_end.strftime('%Y-%m-%d'),
        })
        current = phase_end + timedelta(days=1)

        # Phase 5: Reporting
        phase_end = current + timedelta(days=report_days)
        phases.append({
            'phase': '5',
            'activity': 'SAR Development & Review',
            'start': current.strftime('%Y-%m-%d'),
            'end': phase_end.strftime('%Y-%m-%d'),
        })
        current = phase_end + timedelta(days=1)

        # Phase 6: Final Delivery
        phase_end = current + timedelta(days=3)
        phases.append({
            'phase': '6',
            'activity': 'Final Report Delivery & Briefing',
            'start': current.strftime('%Y-%m-%d'),
            'end': phase_end.strftime('%Y-%m-%d'),
        })

        return phases
