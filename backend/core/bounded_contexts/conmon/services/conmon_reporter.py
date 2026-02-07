"""
Continuous Monitoring (ConMon) Report Generator

Generates comprehensive monthly ConMon reports covering:
- Control validation status and trends
- Vulnerability summary and remediation metrics
- POA&M item status and aging
- Incident tracking summary
- Change management overview
- Evidence freshness analysis
- Executive recommendations

Reports can be exported as Python data structures or Word documents.
"""

import logging
from dataclasses import dataclass, field, asdict
from typing import Dict, List, Optional
from datetime import date, datetime, timedelta
from io import BytesIO

from django.utils import timezone

logger = logging.getLogger(__name__)


@dataclass
class ConMonReport:
    """Comprehensive continuous monitoring report data."""

    period_start: date
    period_end: date
    executive_summary: str = ""
    control_status_summary: Dict = field(default_factory=lambda: {
        "total": 0,
        "validated": 0,
        "at_risk": 0,
        "not_validated": 0,
    })
    vulnerability_summary: Dict = field(default_factory=lambda: {
        "total": 0,
        "critical": 0,
        "high": 0,
        "medium": 0,
        "low": 0,
        "remediated": 0,
    })
    poam_summary: Dict = field(default_factory=lambda: {
        "total": 0,
        "open": 0,
        "overdue": 0,
        "completed": 0,
        "new": 0,
    })
    incident_summary: Dict = field(default_factory=lambda: {
        "total": 0,
        "open": 0,
        "resolved": 0,
    })
    change_summary: Dict = field(default_factory=lambda: {
        "total": 0,
        "approved": 0,
        "pending": 0,
    })
    evidence_freshness: Dict = field(default_factory=lambda: {
        "fresh": 0,
        "stale": 0,
        "missing": 0,
    })
    recommendations: List[str] = field(default_factory=list)
    generated_at: datetime = field(default_factory=timezone.now)

    def to_dict(self) -> Dict:
        result = asdict(self)
        result["period_start"] = self.period_start.isoformat()
        result["period_end"] = self.period_end.isoformat()
        result["generated_at"] = self.generated_at.isoformat()
        return result


class ConMonReporter:
    """
    Generates monthly continuous monitoring reports.

    Aggregates data from multiple sources across the CISO Assistant
    platform to produce a comprehensive ConMon report suitable for
    FedRAMP, NIST, or organizational compliance requirements.
    """

    # Evidence freshness threshold in days
    FRESHNESS_THRESHOLD_DAYS = 90

    def generate_report(
        self,
        period_start: date,
        period_end: date,
        system_id: str = None,
    ) -> ConMonReport:
        """
        Generate a comprehensive ConMon report for the given period.

        Args:
            period_start: Start date of the reporting period.
            period_end: End date of the reporting period.
            system_id: Optional system/project group ID to scope the report.

        Returns:
            A ConMonReport with all sections populated.
        """
        report = ConMonReport(
            period_start=period_start,
            period_end=period_end,
        )

        # Populate each section
        report.control_status_summary = self._gather_control_status(
            period_start, period_end, system_id
        )
        report.vulnerability_summary = self._gather_vulnerability_summary(
            period_start, period_end, system_id
        )
        report.poam_summary = self._gather_poam_summary(
            period_start, period_end, system_id
        )
        report.incident_summary = self._gather_incident_summary(
            period_start, period_end, system_id
        )
        report.change_summary = self._gather_change_summary(
            period_start, period_end, system_id
        )
        report.evidence_freshness = self._gather_evidence_freshness(
            period_end, system_id
        )

        # Generate recommendations and executive summary
        report.recommendations = self._generate_recommendations(report)
        report.executive_summary = self._generate_executive_summary(report)

        logger.info(
            f"ConMon report generated for {period_start} to {period_end}"
        )
        return report

    def export_to_docx(self, report: ConMonReport) -> bytes:
        """
        Export a ConMon report to a Word document (.docx).

        Requires the python-docx package. Returns empty bytes if
        the package is not installed.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning(
                "python-docx is not installed; cannot export to DOCX"
            )
            return b""

        doc = Document()

        # Title
        title = doc.add_heading("Continuous Monitoring Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Period info
        period_para = doc.add_paragraph()
        period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period_run = period_para.add_run(
            f"Reporting Period: {report.period_start} to {report.period_end}"
        )
        period_run.font.size = Pt(12)
        period_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        generated_para = doc.add_paragraph()
        generated_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gen_run = generated_para.add_run(
            f"Generated: {report.generated_at.strftime('%Y-%m-%d %H:%M')}"
        )
        gen_run.font.size = Pt(10)
        gen_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_page_break()

        # Table of Contents placeholder
        doc.add_heading("Table of Contents", level=1)
        toc_items = [
            "1. Executive Summary",
            "2. Control Validation Status",
            "3. Vulnerability Summary",
            "4. POA&M Status",
            "5. Incident Summary",
            "6. Change Management",
            "7. Evidence Freshness",
            "8. Recommendations",
        ]
        for item in toc_items:
            doc.add_paragraph(item, style="List Number")

        doc.add_page_break()

        # 1. Executive Summary
        doc.add_heading("1. Executive Summary", level=1)
        doc.add_paragraph(report.executive_summary or "No summary available.")

        # 2. Control Validation Status
        doc.add_heading("2. Control Validation Status", level=1)
        ctrl = report.control_status_summary
        self._add_summary_table(doc, [
            ("Total Controls", str(ctrl.get("total", 0))),
            ("Validated", str(ctrl.get("validated", 0))),
            ("At Risk", str(ctrl.get("at_risk", 0))),
            ("Not Validated", str(ctrl.get("not_validated", 0))),
        ])

        # 3. Vulnerability Summary
        doc.add_heading("3. Vulnerability Summary", level=1)
        vuln = report.vulnerability_summary
        self._add_summary_table(doc, [
            ("Total Findings", str(vuln.get("total", 0))),
            ("Critical", str(vuln.get("critical", 0))),
            ("High", str(vuln.get("high", 0))),
            ("Medium", str(vuln.get("medium", 0))),
            ("Low", str(vuln.get("low", 0))),
            ("Remediated This Period", str(vuln.get("remediated", 0))),
        ])

        # 4. POA&M Status
        doc.add_heading("4. POA&M Status", level=1)
        poam = report.poam_summary
        self._add_summary_table(doc, [
            ("Total POA&M Items", str(poam.get("total", 0))),
            ("Open", str(poam.get("open", 0))),
            ("Overdue", str(poam.get("overdue", 0))),
            ("Completed This Period", str(poam.get("completed", 0))),
            ("New This Period", str(poam.get("new", 0))),
        ])

        # 5. Incident Summary
        doc.add_heading("5. Incident Summary", level=1)
        inc = report.incident_summary
        self._add_summary_table(doc, [
            ("Total Incidents", str(inc.get("total", 0))),
            ("Open", str(inc.get("open", 0))),
            ("Resolved", str(inc.get("resolved", 0))),
        ])

        # 6. Change Management
        doc.add_heading("6. Change Management", level=1)
        chg = report.change_summary
        self._add_summary_table(doc, [
            ("Total Changes", str(chg.get("total", 0))),
            ("Approved", str(chg.get("approved", 0))),
            ("Pending", str(chg.get("pending", 0))),
        ])

        # 7. Evidence Freshness
        doc.add_heading("7. Evidence Freshness", level=1)
        ev = report.evidence_freshness
        self._add_summary_table(doc, [
            ("Fresh (< 90 days)", str(ev.get("fresh", 0))),
            ("Stale (> 90 days)", str(ev.get("stale", 0))),
            ("Missing", str(ev.get("missing", 0))),
        ])

        # 8. Recommendations
        doc.add_heading("8. Recommendations", level=1)
        if report.recommendations:
            for rec in report.recommendations:
                doc.add_paragraph(rec, style="List Bullet")
        else:
            doc.add_paragraph("No specific recommendations at this time.")

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def get_trend_data(self, months: int = 6) -> List[Dict]:
        """
        Get trend data for ConMon metrics over the specified number of months.

        Returns a list of monthly snapshots with key metrics for trend
        analysis and dashboard visualization.
        """
        trends = []
        now = timezone.now().date()

        for i in range(months - 1, -1, -1):
            # Calculate period boundaries for each month
            if i == 0:
                period_end = now
            else:
                # Go back i months
                year = now.year
                month = now.month - i
                while month <= 0:
                    month += 12
                    year -= 1
                # Last day of that month
                if month == 12:
                    period_end = date(year + 1, 1, 1) - timedelta(days=1)
                else:
                    period_end = date(year, month + 1, 1) - timedelta(days=1)

            # First day of the month for period_end
            period_start = period_end.replace(day=1)

            # Generate a lightweight report for this period
            try:
                report = self.generate_report(period_start, period_end)
                trends.append({
                    "period": period_start.isoformat(),
                    "period_label": period_start.strftime("%b %Y"),
                    "controls_validated": report.control_status_summary.get(
                        "validated", 0
                    ),
                    "controls_at_risk": report.control_status_summary.get(
                        "at_risk", 0
                    ),
                    "vulnerabilities_total": report.vulnerability_summary.get(
                        "total", 0
                    ),
                    "vulnerabilities_remediated": report.vulnerability_summary.get(
                        "remediated", 0
                    ),
                    "poam_open": report.poam_summary.get("open", 0),
                    "poam_overdue": report.poam_summary.get("overdue", 0),
                    "evidence_fresh": report.evidence_freshness.get("fresh", 0),
                    "evidence_stale": report.evidence_freshness.get("stale", 0),
                })
            except Exception as exc:
                logger.warning(
                    f"Failed to generate trend data for {period_start}: {exc}"
                )
                trends.append({
                    "period": period_start.isoformat(),
                    "period_label": period_start.strftime("%b %Y"),
                    "error": str(exc),
                })

        return trends

    # -------------------------------------------------------------------------
    # Private data-gathering methods
    # -------------------------------------------------------------------------

    def _gather_control_status(
        self, period_start: date, period_end: date, system_id: str = None
    ) -> Dict:
        """Gather control validation status metrics."""
        summary = {
            "total": 0,
            "validated": 0,
            "at_risk": 0,
            "not_validated": 0,
        }

        try:
            from core.models import AppliedControl

            qs = AppliedControl.objects.all()
            if system_id:
                qs = qs.filter(folder__content_type__model="project")

            summary["total"] = qs.count()
            summary["validated"] = qs.filter(
                status__in=["active"]
            ).count()
            summary["at_risk"] = qs.filter(
                status__in=["on_hold", "deprecated"]
            ).count()
            summary["not_validated"] = qs.filter(
                status__in=["--", "to_do"]
            ).count()
        except Exception as exc:
            logger.debug(f"Could not gather control status: {exc}")

        return summary

    def _gather_vulnerability_summary(
        self, period_start: date, period_end: date, system_id: str = None
    ) -> Dict:
        """Gather vulnerability findings summary."""
        summary = {
            "total": 0,
            "critical": 0,
            "high": 0,
            "medium": 0,
            "low": 0,
            "remediated": 0,
        }

        # Check ComplianceFinding
        try:
            from core.bounded_contexts.compliance.associations.compliance_finding import (
                ComplianceFinding,
            )

            findings = ComplianceFinding.objects.all()
            summary["total"] = findings.count()
            summary["critical"] = findings.filter(severity="critical").count()
            summary["high"] = findings.filter(severity="high").count()
            summary["medium"] = findings.filter(severity="medium").count()
            summary["low"] = findings.filter(severity="low").count()

            # Remediated during the period
            summary["remediated"] = findings.filter(
                lifecycle_state__in=["verified", "closed"],
                updated_at__gte=period_start,
                updated_at__lte=period_end,
            ).count()
        except Exception as exc:
            logger.debug(f"Could not gather ComplianceFinding data: {exc}")

        # Also check VulnerabilityFinding
        try:
            from core.bounded_contexts.rmf_operations.aggregates.vulnerability_finding import (
                VulnerabilityFinding,
            )

            vuln_findings = VulnerabilityFinding.objects.all()
            summary["total"] += vuln_findings.count()

            # Map severity categories
            for vf in vuln_findings:
                sev = getattr(vf, "severity", "medium")
                if sev in summary:
                    summary[sev] += 1

            # Remediated during period
            summary["remediated"] += vuln_findings.filter(
                status__in=["not_a_finding", "not_applicable"],
                updated_at__gte=period_start,
                updated_at__lte=period_end,
            ).count()
        except Exception as exc:
            logger.debug(f"Could not gather VulnerabilityFinding data: {exc}")

        return summary

    def _gather_poam_summary(
        self, period_start: date, period_end: date, system_id: str = None
    ) -> Dict:
        """Gather POA&M item summary."""
        summary = {
            "total": 0,
            "open": 0,
            "overdue": 0,
            "completed": 0,
            "new": 0,
        }

        try:
            from poam.models.poam_item import POAMItem

            qs = POAMItem.objects.all()
            if system_id:
                qs = qs.filter(system_group_id=system_id)

            summary["total"] = qs.count()

            open_statuses = [
                "draft", "submitted", "approved", "in_progress"
            ]
            summary["open"] = qs.filter(status__in=open_statuses).count()

            # Overdue: past estimated completion date and not completed
            now = timezone.now().date()
            summary["overdue"] = qs.filter(
                estimated_completion_date__lt=now,
                status__in=open_statuses,
            ).count()

            # Completed during the reporting period
            summary["completed"] = qs.filter(
                status="completed",
                actual_completion_date__gte=period_start,
                actual_completion_date__lte=period_end,
            ).count()

            # New during the reporting period
            summary["new"] = qs.filter(
                identified_date__gte=period_start,
                identified_date__lte=period_end,
            ).count()
        except Exception as exc:
            logger.debug(f"Could not gather POA&M data: {exc}")

        return summary

    def _gather_incident_summary(
        self, period_start: date, period_end: date, system_id: str = None
    ) -> Dict:
        """Gather security incident summary."""
        summary = {
            "total": 0,
            "open": 0,
            "resolved": 0,
        }

        # Try business continuity incidents if available
        try:
            from core.bounded_contexts.security_operations.aggregates import (
                SecurityIncident,
            )

            incidents = SecurityIncident.objects.filter(
                created_at__gte=period_start,
                created_at__lte=period_end,
            )
            summary["total"] = incidents.count()
            summary["open"] = incidents.filter(
                status__in=["open", "investigating"]
            ).count()
            summary["resolved"] = incidents.filter(
                status__in=["resolved", "closed"]
            ).count()
        except Exception:
            pass

        # Fallback: try the core incident models if they exist
        try:
            from core.models import Incident

            incidents = Incident.objects.filter(
                created_at__gte=period_start,
                created_at__lte=period_end,
            )
            summary["total"] = incidents.count()
        except Exception:
            pass

        return summary

    def _gather_change_summary(
        self, period_start: date, period_end: date, system_id: str = None
    ) -> Dict:
        """Gather change management summary."""
        summary = {
            "total": 0,
            "approved": 0,
            "pending": 0,
        }

        # Try to get change data from audit logs or version history
        try:
            from auditlog.models import LogEntry

            entries = LogEntry.objects.filter(
                timestamp__gte=period_start,
                timestamp__lte=period_end,
                action=1,  # UPDATE
            )
            summary["total"] = entries.count()
            summary["approved"] = entries.count()  # All logged changes are applied
        except Exception:
            pass

        return summary

    def _gather_evidence_freshness(
        self, as_of_date: date, system_id: str = None
    ) -> Dict:
        """Gather evidence freshness metrics."""
        freshness = {
            "fresh": 0,
            "stale": 0,
            "missing": 0,
        }

        try:
            from core.models import AppliedControl, Evidence

            freshness_cutoff = as_of_date - timedelta(
                days=self.FRESHNESS_THRESHOLD_DAYS
            )

            controls = AppliedControl.objects.prefetch_related("evidences")
            if system_id:
                controls = controls.filter(
                    folder__content_type__model="project"
                )

            for control in controls:
                evidences = control.evidences.all()
                if not evidences.exists():
                    freshness["missing"] += 1
                    continue

                has_fresh = False
                for ev in evidences:
                    last_rev = getattr(ev, "last_revision", None)
                    if last_rev and hasattr(last_rev, "created_at"):
                        ev_date = last_rev.created_at.date() if hasattr(
                            last_rev.created_at, "date"
                        ) else last_rev.created_at
                        if ev_date >= freshness_cutoff:
                            has_fresh = True
                            break
                    elif hasattr(ev, "updated_at") and ev.updated_at:
                        ev_date = ev.updated_at.date() if hasattr(
                            ev.updated_at, "date"
                        ) else ev.updated_at
                        if ev_date >= freshness_cutoff:
                            has_fresh = True
                            break

                if has_fresh:
                    freshness["fresh"] += 1
                else:
                    freshness["stale"] += 1
        except Exception as exc:
            logger.debug(f"Could not gather evidence freshness: {exc}")

        return freshness

    def _generate_recommendations(self, report: ConMonReport) -> List[str]:
        """Generate automated recommendations based on report data."""
        recommendations = []

        # POA&M overdue items
        overdue = report.poam_summary.get("overdue", 0)
        if overdue > 0:
            recommendations.append(
                f"Address {overdue} overdue POA&M item(s). Overdue items "
                f"increase organizational risk and may affect authorization "
                f"status."
            )

        # Critical vulnerabilities
        critical = report.vulnerability_summary.get("critical", 0)
        if critical > 0:
            recommendations.append(
                f"Prioritize remediation of {critical} critical "
                f"vulnerability finding(s). Critical findings should be "
                f"addressed within 30 days per FedRAMP requirements."
            )

        # High vulnerabilities
        high = report.vulnerability_summary.get("high", 0)
        if high > 0:
            recommendations.append(
                f"Schedule remediation for {high} high severity "
                f"finding(s). High findings should be addressed within "
                f"90 days."
            )

        # Stale evidence
        stale = report.evidence_freshness.get("stale", 0)
        if stale > 0:
            recommendations.append(
                f"Update evidence for {stale} control(s) with stale "
                f"evidence (older than {self.FRESHNESS_THRESHOLD_DAYS} days). "
                f"Fresh evidence is required for continuous monitoring."
            )

        # Missing evidence
        missing = report.evidence_freshness.get("missing", 0)
        if missing > 0:
            recommendations.append(
                f"Collect evidence for {missing} control(s) currently "
                f"lacking any supporting evidence."
            )

        # Controls at risk
        at_risk = report.control_status_summary.get("at_risk", 0)
        if at_risk > 0:
            recommendations.append(
                f"Review {at_risk} control(s) currently flagged as at-risk. "
                f"These controls may require additional compensating measures."
            )

        # No issues
        if not recommendations:
            recommendations.append(
                "All continuous monitoring metrics are within acceptable "
                "thresholds. Continue regular monitoring activities."
            )

        return recommendations

    def _generate_executive_summary(self, report: ConMonReport) -> str:
        """Generate an executive summary narrative."""
        ctrl = report.control_status_summary
        vuln = report.vulnerability_summary
        poam = report.poam_summary
        ev = report.evidence_freshness

        total_controls = ctrl.get("total", 0)
        validated = ctrl.get("validated", 0)
        pct_validated = (
            round((validated / total_controls) * 100)
            if total_controls > 0
            else 0
        )

        total_vulns = vuln.get("total", 0)
        remediated = vuln.get("remediated", 0)
        critical = vuln.get("critical", 0)

        poam_open = poam.get("open", 0)
        poam_overdue = poam.get("overdue", 0)
        poam_completed = poam.get("completed", 0)

        fresh = ev.get("fresh", 0)
        stale = ev.get("stale", 0)
        missing = ev.get("missing", 0)
        total_evidence_controls = fresh + stale + missing
        pct_fresh = (
            round((fresh / total_evidence_controls) * 100)
            if total_evidence_controls > 0
            else 0
        )

        summary_parts = [
            f"During the reporting period ({report.period_start} to "
            f"{report.period_end}), the continuous monitoring program "
            f"assessed {total_controls} controls with {pct_validated}% "
            f"validated as active.",
        ]

        if total_vulns > 0:
            summary_parts.append(
                f" A total of {total_vulns} vulnerability findings were "
                f"tracked, with {remediated} remediated during the period."
            )
            if critical > 0:
                summary_parts.append(
                    f" {critical} critical finding(s) require immediate "
                    f"attention."
                )

        if poam_open > 0:
            summary_parts.append(
                f" The POA&M register contains {poam_open} open items"
            )
            if poam_overdue > 0:
                summary_parts.append(
                    f", of which {poam_overdue} are overdue"
                )
            summary_parts.append(
                f". {poam_completed} items were completed during the period."
            )

        summary_parts.append(
            f" Evidence freshness stands at {pct_fresh}%, with {stale} "
            f"controls having stale evidence and {missing} lacking "
            f"evidence entirely."
        )

        return "".join(summary_parts)

    # -------------------------------------------------------------------------
    # Helper for docx export
    # -------------------------------------------------------------------------

    def _add_summary_table(self, doc, rows: List[tuple]):
        """Add a two-column summary table to the Word document."""
        try:
            from docx.shared import Pt, RGBColor, Inches

            table = doc.add_table(rows=len(rows), cols=2)
            table.style = "Light Grid Accent 1"

            for i, (label, value) in enumerate(rows):
                cell_label = table.cell(i, 0)
                cell_value = table.cell(i, 1)
                cell_label.text = label
                cell_value.text = value

                # Style the value cell
                for paragraph in cell_value.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            doc.add_paragraph("")  # Spacer
        except Exception as exc:
            # Fallback: just add as paragraphs
            for label, value in rows:
                doc.add_paragraph(f"{label}: {value}")


def get_conmon_reporter() -> ConMonReporter:
    """Factory function to create a ConMonReporter instance."""
    return ConMonReporter()
